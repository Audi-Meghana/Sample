"""
Multi-format document text extractor for gene analysis
Supports: DOCX, DOC, TXT, XLSX, RTF, and other text-based documents
"""

import io
from typing import Optional


def extract_text_from_document(file_obj, filename: str = "") -> str:
    """
    Extract text from various document formats
    
    Args:
        file_obj: File-like object or bytes
        filename: Original filename (optional, helps with format detection)
    
    Returns:
        Extracted text or empty string if extraction fails
    """
    
    text = ""
    
    try:
        # Convert file_obj to bytes if it's a file-like object
        if hasattr(file_obj, 'read'):
            file_bytes = file_obj.read()
            if hasattr(file_obj, 'seek'):
                file_obj.seek(0)
        else:
            file_bytes = file_obj
        
        # Detect format from filename or content
        filename_lower = filename.lower() if filename else ""
        
        # ===== WORD DOCUMENTS (.docx, .doc) =====
        if filename_lower.endswith(('.docx', '.doc')) or 'word' in filename_lower:
            text = _extract_docx(file_bytes) or _extract_doc(file_bytes) or ""
        
        # ===== PLAIN TEXT (.txt) =====
        elif filename_lower.endswith('.txt'):
            try:
                text = file_bytes.decode('utf-8', errors='replace').strip()
            except:
                text = file_bytes.decode('latin-1', errors='replace').strip()
        
        # ===== EXCEL/SPREADSHEETS (.xlsx, .xls) =====
        elif filename_lower.endswith(('.xlsx', '.xls', '.csv')) or 'sheet' in filename_lower or 'excel' in filename_lower:
            text = _extract_excel(file_bytes) or ""
        
        # ===== RTF DOCUMENTS =====
        elif filename_lower.endswith('.rtf'):
            text = _extract_rtf(file_bytes) or ""
        
        # ===== FALLBACK: Try all extractors =====
        else:
            for extractor in [_extract_docx, _extract_doc, _extract_excel, _extract_rtf]:
                text = extractor(file_bytes)
                if text and len(text.strip()) > 20:
                    break
            
            # Last resort: treat as text
            if not text or len(text.strip()) < 20:
                try:
                    text = file_bytes.decode('utf-8', errors='replace').strip()
                except:
                    text = file_bytes.decode('latin-1', errors='replace').strip()
    
    except Exception as e:
        print(f"Error extracting text from document: {str(e)}")
        return ""
    
    return text.strip() if text else ""


def _extract_docx(file_bytes: bytes) -> Optional[str]:
    """Extract text from DOCX (Office Open XML) files"""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        return text
    except:
        return None


def _extract_doc(file_bytes: bytes) -> Optional[str]:
    """Extract text from older DOC files"""
    try:
        from docx2docx import convert
        # Try using python-docx with fallback
        # Most .doc files can be read if they're more recent
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        return text
    except:
        # Alternative: try using python-pptx or other methods
        try:
            import zipfile
            # Some DOC files are actually ZIP archives
            with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
                if 'document.xml' in zf.namelist():
                    with zf.open('document.xml') as xml_file:
                        xml_content = xml_file.read()
                        # Extract text from XML (crude but effective)
                        import re
                        text = re.findall(r'<w:t>(.*?)</w:t>', xml_content.decode('utf-8', errors='ignore'))
                        return "\n".join(text)
        except:
            pass
        return None


def _extract_excel(file_bytes: bytes) -> Optional[str]:
    """Extract text from Excel/CSV files"""
    try:
        import pandas as pd
        
        # Try XLSX first
        try:
            df = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None)
            text = ""
            for sheet_name, sheet_df in df.items():
                text += f"Sheet: {sheet_name}\n"
                text += sheet_df.to_string()
                text += "\n\n"
            return text
        except:
            pass
        
        # Try CSV
        try:
            df = pd.read_csv(io.BytesIO(file_bytes))
            return df.to_string()
        except:
            pass
        
        return None
    except:
        return None


def _extract_rtf(file_bytes: bytes) -> Optional[str]:
    """Extract text from RTF files"""
    try:
        from striprtf.striprtf import rtf_to_text
        text = rtf_to_text(file_bytes.decode('utf-8', errors='replace'))
        return text
    except:
        # Fallback: basic RTF parsing
        try:
            rtf_text = file_bytes.decode('utf-8', errors='replace')
            # Remove RTF control sequences
            import re
            text = re.sub(r'\\[a-z]+\d*\s?', ' ', rtf_text)
            text = re.sub(r'[{}\\]', '', text)
            return text.strip()
        except:
            return None
